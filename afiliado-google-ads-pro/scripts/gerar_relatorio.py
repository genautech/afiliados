#!/usr/bin/env python3
"""Gera relatório de performance (.docx) a partir de um CSV de campanhas.

CSV esperado (cabeçalho flexível, separador , ou ;):
  campanha, gasto, impressoes, cliques, conversoes
Colunas extras são ignoradas. Aceita export do Google Ads renomeando colunas.

Uso:
  python3 gerar_relatorio.py dados.csv --comissao 97 [--roas-meta 2.0] \
      [--moeda BRL] [--periodo "01–07/07/2026"] -o relatorio.docx
"""
import argparse
import csv
import io
import unicodedata
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0xE0, 0x7B, 0x39)
SIMB = {"BRL": "R$", "USD": "US$", "GBP": "£", "AUD": "A$"}

ALIASES = {
    "campanha": ["campanha", "campaign", "campaign name", "nome da campanha"],
    "gasto": ["gasto", "custo", "cost", "spend", "custo (brl)"],
    "impressoes": ["impressoes", "impressões", "impressions", "impr."],
    "cliques": ["cliques", "clicks"],
    "conversoes": ["conversoes", "conversões", "conversions", "conv."],
}


def _norm(s):
    s = unicodedata.normalize("NFKD", s.lower().strip())
    return "".join(c for c in s if not unicodedata.combining(c))


def _num(v):
    v = str(v).strip().replace("R$", "").replace("US$", "").replace("£", "").replace("A$", "").strip()
    if v.count(",") == 1 and (v.count(".") == 0 or v.rfind(",") > v.rfind(".")):
        v = v.replace(".", "").replace(",", ".")
    else:
        v = v.replace(",", "")
    try:
        return float(v)
    except ValueError:
        return 0.0


def ler_csv(path):
    raw = open(path, encoding="utf-8-sig").read()
    delim = ";" if raw.split("\n")[0].count(";") > raw.split("\n")[0].count(",") else ","
    rd = csv.DictReader(io.StringIO(raw), delimiter=delim)
    colmap = {}
    for campo, opts in ALIASES.items():
        for h in rd.fieldnames or []:
            if _norm(h) in [_norm(o) for o in opts]:
                colmap[campo] = h
                break
    linhas = []
    for row in rd:
        nome = row.get(colmap.get("campanha", ""), "").strip()
        if not nome:
            continue
        linhas.append({
            "campanha": nome,
            "gasto": _num(row.get(colmap.get("gasto", ""), 0)),
            "impressoes": _num(row.get(colmap.get("impressoes", ""), 0)),
            "cliques": _num(row.get(colmap.get("cliques", ""), 0)),
            "conversoes": _num(row.get(colmap.get("conversoes", ""), 0)),
        })
    return linhas


def veredicto(l, comissao, roas_meta):
    receita = l["conversoes"] * comissao
    roas = receita / l["gasto"] if l["gasto"] else 0
    cpa = l["gasto"] / l["conversoes"] if l["conversoes"] else None
    amostra = l["cliques"] >= 100
    if l["conversoes"] == 0 and l["gasto"] >= 3 * comissao:
        return "✖ PAUSAR / TROCAR OFERTA", f"Regra do 3×: gastou {l['gasto']:.2f} (≥3× a comissão) sem venda. Confirme o rastreamento; se OK, troque a oferta antes do nicho."
    if not amostra:
        return "⏳ AGUARDAR AMOSTRA", f"Apenas {int(l['cliques'])} cliques — mínimo ~100/keyword ou 7–14 dias antes de decidir. Enquanto isso, negativar termos de pesquisa irrelevantes."
    if l["conversoes"] == 0:
        return "⚙ INVESTIGAR", "Amostra razoável e 0 conversões: checar rastreamento, velocidade da página-ponte e intenção das keywords."
    if roas >= roas_meta:
        return "✔ ESCALAR", f"ROAS {roas:.2f} ≥ meta {roas_meta}. Aumentar orçamento +20–30% e reavaliar em 4–7 dias."
    if roas >= 1:
        return "⚙ OTIMIZAR", f"Lucrativa (ROAS {roas:.2f}) mas abaixo da meta. Negativar desperdício e pausar keywords sem conversão com >30 cliques — UMA mudança por vez."
    return "✖ CORTAR DESPERDÍCIO", f"ROAS {roas:.2f} < 1 (CPA {cpa:.2f} acima do breakeven). Analisar dispositivo/horário/região antes de pausar."


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("csv")
    ap.add_argument("--comissao", type=float, required=True)
    ap.add_argument("--roas-meta", type=float, default=2.0)
    ap.add_argument("--moeda", default="BRL")
    ap.add_argument("--periodo", default="período informado")
    ap.add_argument("-o", "--out", default="relatorio-performance.docx")
    a = ap.parse_args()
    simb = SIMB.get(a.moeda, a.moeda)

    linhas = ler_csv(a.csv)
    if not linhas:
        raise SystemExit("Nenhuma linha válida no CSV — confira o cabeçalho (campanha, gasto, impressoes, cliques, conversoes).")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RELATÓRIO DE PERFORMANCE — GOOGLE ADS"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(f"Período: {a.periodo} · Comissão média: {simb} {a.comissao:.2f} · ROAS-meta: {a.roas_meta}")
    r.italic = True; r.font.color.rgb = ACCENT

    tot = {k: sum(l[k] for l in linhas) for k in ("gasto", "impressoes", "cliques", "conversoes")}
    receita = tot["conversoes"] * a.comissao
    roas = receita / tot["gasto"] if tot["gasto"] else 0
    h = doc.add_paragraph(); r = h.add_run("Resumo executivo"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    doc.add_paragraph(
        f"Investimento total de {simb} {tot['gasto']:.2f} gerou {int(tot['cliques'])} cliques "
        f"({int(tot['impressoes'])} impressões) e {int(tot['conversoes'])} conversões — receita estimada de "
        f"{simb} {receita:.2f} e ROAS geral de {roas:.2f} "
        f"({'acima' if roas >= a.roas_meta else 'abaixo'} da meta de {a.roas_meta}). "
        f"Vereditos por campanha e plano de ação abaixo. Valores de receita usam a comissão média informada; "
        f"confirme na plataforma."
    )

    h = doc.add_paragraph(); r = h.add_run("Campanhas"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    tab = doc.add_table(rows=1, cols=8); tab.style = "Light Grid Accent 1"
    for i, c in enumerate(["Campanha", f"Gasto ({simb})", "Cliques", "CTR", f"CPC ({simb})", "Conv.", f"CPA ({simb})", "ROAS"]):
        cell = tab.rows[0].cells[i]; cell.text = c
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for l in linhas:
        ctr = l["cliques"] / l["impressoes"] * 100 if l["impressoes"] else 0
        cpc = l["gasto"] / l["cliques"] if l["cliques"] else 0
        cpa = l["gasto"] / l["conversoes"] if l["conversoes"] else None
        rroas = (l["conversoes"] * a.comissao) / l["gasto"] if l["gasto"] else 0
        row = tab.add_row().cells
        vals = [l["campanha"], f"{l['gasto']:.2f}", f"{int(l['cliques'])}", f"{ctr:.2f}%",
                f"{cpc:.2f}", f"{int(l['conversoes'])}", f"{cpa:.2f}" if cpa else "—", f"{rroas:.2f}"]
        for i, v in enumerate(vals):
            row[i].text = str(v)

    h = doc.add_paragraph(); r = h.add_run("Veredictos e plano de ação"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    for l in linhas:
        v, txt = veredicto(l, a.comissao, a.roas_meta)
        p = doc.add_paragraph()
        r = p.add_run(f"{l['campanha']}: {v}. "); r.bold = True
        p.add_run(txt)

    doc.add_paragraph().add_run(
        "Lembretes do sistema: uma mudança por vez; decisões exigem 7–14 dias ou ~100 cliques; "
        "escala é +20–30% por vez; sem rastreamento comprovado, nenhum veredicto é confiável."
    ).italic = True

    doc.save(a.out)
    print(f"Relatório gerado: {a.out} ({len(linhas)} campanhas)")


if __name__ == "__main__":
    main()
